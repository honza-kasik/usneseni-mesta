#!/usr/bin/env python3
"""Extract text from downloaded Litovel ZM archive documents."""

from __future__ import annotations

import argparse
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

import pdfplumber

try:
    from archive_zm_common import read_json, write_json
except ImportError:
    from tools.archive_zm_common import read_json, write_json


def normalize_text(text: str) -> str:
    import re

    text = text.replace("\xa0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def probably_binary_garbage(text: str) -> bool:
    if not text:
        return False
    replacement_chars = text.count("\ufffd")
    control_chars = sum(1 for char in text if ord(char) < 32 and char not in "\n\r\t")
    length = len(text)
    return (replacement_chars / length) > 0.01 or (control_chars / length) > 0.05


def text_quality_flag(text: str, error: str | None, short_text_threshold: int) -> str:
    if error:
        return "extraction_failed"
    if not text:
        return "empty_text"
    if probably_binary_garbage(text):
        return "probably_binary_garbage"
    if len(text) < short_text_threshold:
        return "short_text"
    return "text_ok"


def extract_pdf(path: Path) -> tuple[str, str, int | None]:
    pages = []
    with pdfplumber.open(path) as pdf:
        page_count = len(pdf.pages)
        for page in pdf.pages:
            pages.append(page.extract_text() or "")
    return normalize_text("\n".join(pages)), "pdf_text", page_count


def extract_docx(path: Path) -> tuple[str, str, int | None]:
    from docx import Document

    document = Document(path)
    paragraphs = [p.text for p in document.paragraphs if p.text]
    for table in document.tables:
        for row in table.rows:
            paragraphs.append(" ".join(cell.text for cell in row.cells if cell.text))
    return normalize_text("\n".join(paragraphs)), "docx", None


def extract_doc_with_libreoffice(path: Path) -> tuple[str, str, int | None]:
    executable = shutil.which("libreoffice") or shutil.which("soffice")
    if not executable:
        raise FileNotFoundError("libreoffice/soffice not found")

    with tempfile.TemporaryDirectory() as tmp:
        outdir = Path(tmp)
        subprocess.run(
            [
                executable,
                "--headless",
                "--convert-to",
                "txt:Text",
                "--outdir",
                str(outdir),
                str(path),
            ],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            timeout=60,
        )
        converted = outdir / f"{path.stem}.txt"
        if not converted.exists():
            candidates = list(outdir.glob("*.txt"))
            if not candidates:
                raise FileNotFoundError("LibreOffice did not produce a .txt file")
            converted = candidates[0]
        return normalize_text(converted.read_text(encoding="utf-8", errors="replace")), "libreoffice", None


def extract_doc_with_command(path: Path, command: str) -> tuple[str, str, int | None]:
    executable = shutil.which(command)
    if not executable:
        raise FileNotFoundError(f"{command} not found")
    result = subprocess.run(
        [executable, str(path)],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        timeout=60,
    )
    return normalize_text(result.stdout.decode("utf-8", errors="replace")), command, None


def extract_doc(path: Path) -> tuple[str, str, int | None]:
    errors = []
    for extractor in (
        extract_doc_with_libreoffice,
        lambda p: extract_doc_with_command(p, "antiword"),
        lambda p: extract_doc_with_command(p, "catdoc"),
    ):
        try:
            return extractor(path)
        except Exception as exc:
            errors.append(str(exc))
    raise RuntimeError("; ".join(errors))


def extract_file(path: Path, file_type: str) -> tuple[str, str, int | None]:
    if file_type == "pdf":
        return extract_pdf(path)
    if file_type == "docx":
        return extract_docx(path)
    if file_type == "doc":
        return extract_doc(path)
    raise ValueError(f"Unsupported file type: {file_type}")


def extract_inventory(
    inventory: list[dict],
    text_dir: Path,
    short_text_threshold: int = 200,
) -> tuple[list[dict], dict]:
    text_dir.mkdir(parents=True, exist_ok=True)
    extraction: dict[str, dict] = {}

    for item in inventory:
        item_id = item["id"]
        local_path = item.get("local_path")
        file_type = item.get("file_type")
        if item.get("status") != "downloaded" or not local_path:
            extraction[item_id] = {
                "method": "none",
                "text_path": None,
                "text_chars": 0,
                "has_text": False,
                "quality_flag": "extraction_failed",
                "error": "not_downloaded",
            }
            continue

        try:
            text, method, page_count = extract_file(Path(local_path), file_type)
            text_path = text_dir / f"{item_id}.txt"
            text_path.write_text(text + ("\n" if text else ""), encoding="utf-8")
            text_chars = len(text)
            chars_per_page = (text_chars / page_count) if page_count else None
            extraction[item_id] = {
                "method": method,
                "text_path": str(text_path),
                "text_chars": text_chars,
                "chars_per_page": chars_per_page,
                "has_text": bool(text),
                "quality_flag": text_quality_flag(text, None, short_text_threshold),
                "error": None,
            }
        except Exception as exc:
            extraction[item_id] = {
                "method": "none",
                "text_path": None,
                "text_chars": 0,
                "chars_per_page": None,
                "has_text": False,
                "quality_flag": "extraction_failed",
                "error": str(exc),
            }

    return inventory, extraction


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--inventory", type=Path, default=Path("work/archive_zm/inventory.json"))
    parser.add_argument("--text-dir", type=Path, default=Path("work/archive_zm/text"))
    parser.add_argument("--output", type=Path, default=Path("work/archive_zm/extraction.json"))
    parser.add_argument("--short-text-threshold", type=int, default=200)
    args = parser.parse_args()

    inventory = read_json(args.inventory, [])
    if not isinstance(inventory, list):
        raise SystemExit("Inventory must be a JSON list.")

    _, extraction = extract_inventory(inventory, args.text_dir, args.short_text_threshold)
    write_json(args.output, extraction)
    print(f"Extraction metadata written: {args.output}")
    print(f"With text: {sum(1 for item in extraction.values() if item.get('has_text'))}")
    print(f"Without text: {sum(1 for item in extraction.values() if not item.get('has_text'))}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
